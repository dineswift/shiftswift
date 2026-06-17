"""Export HR template markdown as PDF or Word (.docx)."""

from __future__ import annotations

import html
import io
import re
from datetime import datetime, timezone


def _inline_markup(text: str) -> str:
    escaped = html.escape(text)
    escaped = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", escaped)
    escaped = re.sub(r"\*(.+?)\*", r"<i>\1</i>", escaped)
    return escaped


def _add_docx_rich_paragraph(document, text: str, *, style: str | None = None, italic: bool = False) -> None:
    paragraph = document.add_paragraph(style=style)
    parts = re.split(r"(\*\*.+?\*\*|\*.+?\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            if italic:
                run.italic = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(part)
            if italic:
                run.italic = True


def _markdown_to_docx(document, markdown: str) -> None:
    bullet_items: list[str] = []

    def flush_bullets() -> None:
        nonlocal bullet_items
        for item in bullet_items:
            _add_docx_rich_paragraph(document, item, style="List Bullet")
        bullet_items = []

    for line in markdown.splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            flush_bullets()
            document.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            flush_bullets()
            document.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            flush_bullets()
            document.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- "):
            bullet_items.append(stripped[2:])
        elif stripped.startswith("- [ ] ") or stripped.startswith("- [x] "):
            prefix = "☐ " if stripped.startswith("- [ ] ") else "☑ "
            bullet_items.append(prefix + stripped[6:])
        elif stripped.startswith("|") and "---" in stripped:
            continue
        elif stripped.startswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            flush_bullets()
            _add_docx_rich_paragraph(document, " · ".join(cells))
        elif stripped == "---":
            flush_bullets()
            document.add_paragraph()
        elif not stripped:
            flush_bullets()
            document.add_paragraph()
        else:
            flush_bullets()
            if stripped.startswith("_") and stripped.endswith("_"):
                _add_docx_rich_paragraph(document, stripped.strip("_"), italic=True)
            else:
                _add_docx_rich_paragraph(document, stripped)

    flush_bullets()


def _markdown_to_pdf_flowables(markdown: str) -> list:
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import ListFlowable, ListItem, Paragraph, Spacer

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle(
        "TemplateH1",
        parent=styles["Heading1"],
        fontSize=16,
        spaceAfter=8,
        textColor=colors.HexColor("#0F6E56"),
    )
    h2 = ParagraphStyle(
        "TemplateH2",
        parent=styles["Heading2"],
        fontSize=13,
        spaceBefore=10,
        spaceAfter=6,
        textColor=colors.HexColor("#1E293B"),
    )
    h3 = ParagraphStyle(
        "TemplateH3",
        parent=styles["Heading3"],
        fontSize=11,
        spaceBefore=8,
        spaceAfter=4,
        textColor=colors.HexColor("#334155"),
    )
    body = ParagraphStyle("TemplateBody", parent=styles["Normal"], fontSize=10, leading=14)
    italic = ParagraphStyle("TemplateItalic", parent=body, fontName="Helvetica-Oblique")

    flowables: list = []
    bullet_items: list[str] = []

    def flush_bullets() -> None:
        nonlocal bullet_items
        if not bullet_items:
            return
        items = [ListItem(Paragraph(_inline_markup(item), body), leftIndent=12) for item in bullet_items]
        flowables.append(ListFlowable(items, bulletType="bullet", start="•", leftIndent=18))
        bullet_items = []

    for line in markdown.splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            flush_bullets()
            flowables.append(Paragraph(_inline_markup(stripped[2:]), h1))
        elif stripped.startswith("## "):
            flush_bullets()
            flowables.append(Paragraph(_inline_markup(stripped[3:]), h2))
        elif stripped.startswith("### "):
            flush_bullets()
            flowables.append(Paragraph(_inline_markup(stripped[4:]), h3))
        elif stripped.startswith("- "):
            bullet_items.append(stripped[2:])
        elif stripped.startswith("- [ ] ") or stripped.startswith("- [x] "):
            prefix = "☐ " if stripped.startswith("- [ ] ") else "☑ "
            text = stripped[6:]
            bullet_items.append(prefix + text)
        elif stripped.startswith("|") and "---" in stripped:
            continue
        elif stripped.startswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            flush_bullets()
            flowables.append(Paragraph(_inline_markup(" · ".join(cells)), body))
        elif stripped == "---":
            flush_bullets()
            flowables.append(Spacer(1, 6))
        elif not stripped:
            flush_bullets()
            flowables.append(Spacer(1, 4))
        else:
            flush_bullets()
            if stripped.startswith("_") and stripped.endswith("_"):
                flowables.append(Paragraph(_inline_markup(stripped.strip("_")), italic))
            else:
                flowables.append(Paragraph(_inline_markup(stripped), body))

    flush_bullets()
    return flowables


def build_template_pdf_bytes(markdown: str, *, title: str) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title=title,
    )
    styles = getSampleStyleSheet()
    footer_style = ParagraphStyle(
        "TemplateFooter",
        parent=styles["Normal"],
        fontSize=8,
        textColor=colors.HexColor("#94A3B8"),
    )
    generated = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    body: list = [
        Paragraph(
            f"<i>Generated by ShiftSwift HR (shiftswifthr.co.uk) · {generated}</i>",
            footer_style,
        ),
        Spacer(1, 8),
    ]
    body.extend(_markdown_to_pdf_flowables(markdown))
    doc.build(body)
    return buffer.getvalue()


def build_template_docx_bytes(markdown: str, *, title: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor

    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    generated = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    meta = document.add_paragraph()
    meta_run = meta.add_run(f"Generated by ShiftSwift HR (shiftswifthr.co.uk) · {generated}")
    meta_run.italic = True
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    document.add_paragraph()

    _markdown_to_docx(document, markdown)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_template_word_bytes(markdown: str, *, title: str) -> bytes:
    """Word download — real Office Open XML (.docx), not HTML masquerading as .doc."""
    return build_template_docx_bytes(markdown, title=title)
